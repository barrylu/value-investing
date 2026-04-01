#!/usr/bin/env python3
"""
巴菲特致股东信+股东大会文档拆分脚本
将19MB/440万字的巨型文档拆分为独立的Markdown文件
"""

import re
import os
from docx import Document

DOC_PATH = "/Users/luzuoguan/ai/value-investing/文档/巴菲特1950-2022致股东信+股东大会_20230626220808.docx"
OUTPUT_BASE = "/Users/luzuoguan/WorkBuddy/Claw/value-investing/巴菲特文集"

def get_text_with_tables(doc, start_idx, end_idx):
    """Extract text from paragraphs, preserving table content inline."""
    lines = []
    for i in range(start_idx, end_idx):
        if i >= len(doc.paragraphs):
            break
        text = doc.paragraphs[i].text.strip()
        if text:
            lines.append(text)
        elif lines and lines[-1] != '':
            lines.append('')  # preserve paragraph breaks
    # Clean trailing empty lines
    while lines and lines[-1] == '':
        lines.pop()
    return '\n\n'.join(line for line in lines if line != '') if lines else ''


def extract_tables_near(doc, para_idx_start, para_idx_end):
    """Extract tables that fall within the paragraph range.
    Since python-docx doesn't easily map tables to paragraph positions,
    we'll handle tables separately."""
    # Tables will be extracted in a second pass
    pass


def find_sections(doc):
    """Identify all major sections in the document."""
    sections = []
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        
        # 致股东信 headers: "巴菲特致股东的信 YYYY" / "巴菲特给股东的信 YYYY"
        m = re.match(r'^巴菲特[致给]股东的信\s*(\d{4})', text)
        if m:
            sections.append({
                'type': '致股东信',
                'year': int(m.group(1)),
                'para_idx': i,
                'title': text
            })
            continue
        
        # 股东大会 headers: "伯克希尔股东大会实录 YYYY" or "伯克希尔股东大会问答 YYYY"
        m = re.match(r'^伯克希尔股东大会(?:实录|问答)\s*(\d{4})', text)
        if m:
            sections.append({
                'type': '股东大会',
                'year': int(m.group(1)),
                'para_idx': i,
                'title': text
            })
            continue
        
        # BH early shareholder letters: "伯克希尔.哈撒韦公司 YYYY年"
        m = re.match(r'^伯克希尔[.·]哈撒韦公司\s*(\d{4})\s*年', text)
        if m:
            sections.append({
                'type': '致股东信',
                'year': int(m.group(1)),
                'para_idx': i,
                'title': text
            })
            continue
    
    return sections


def find_partnership_letters(doc):
    """Find Buffett Partnership letters (1957-1969)."""
    sections = []
    partnership_mode = False
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        
        # Partnership letters have date headers like "1957年1月", "1959年2月11日"
        # They appear between roughly para 234 and para 3400
        if i < 200 or i > 3700:
            continue
            
        # Match year+month date headers that start sections
        m = re.match(r'^(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})?\s*日?$', text)
        if m:
            year = int(m.group(1))
            if 1957 <= year <= 1970:
                sections.append({
                    'type': '合伙人信',
                    'year': year,
                    'para_idx': i,
                    'title': f'巴菲特合伙人信 {text}'
                })
                continue
    
    return sections


def find_early_articles(doc):
    """Find early articles (pre-partnership)."""
    sections = []
    
    # "我最看好的股票：GEICO保险 1951" around para 129
    # "西部保险" article around para 183
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text or i > 250:
            continue
        
        if '我最看好的股票' in text and 'GEICO' in text:
            sections.append({
                'type': '早期文章',
                'year': 1951,
                'para_idx': i,
                'title': '我最看好的股票：GEICO保险 (1951)'
            })
        elif i >= 183 and i <= 185 and ('1953' in text or '西部保险' in text):
            sections.append({
                'type': '早期文章',
                'year': 1953,
                'para_idx': i,
                'title': '西部保险公司分析 (1953)'
            })
    
    return sections


def merge_and_sort_sections(all_sections):
    """Merge all sections, sort by paragraph index, deduplicate."""
    # Sort by paragraph index
    all_sections.sort(key=lambda x: x['para_idx'])
    
    # Remove duplicates (same year + type with very close para_idx)
    deduped = []
    for s in all_sections:
        if deduped and deduped[-1]['year'] == s['year'] and deduped[-1]['type'] == s['type']:
            # Keep the one with earlier para_idx
            if s['para_idx'] - deduped[-1]['para_idx'] < 50:
                continue
        deduped.append(s)
    
    return deduped


def write_section(doc, section, next_para_idx, output_base):
    """Write a section to a markdown file."""
    type_dir = section['type']
    year = section['year']
    start = section['para_idx']
    end = next_para_idx
    
    # Build output path
    dir_path = os.path.join(output_base, type_dir)
    os.makedirs(dir_path, exist_ok=True)
    
    # Filename
    if section['type'] == '合伙人信':
        # Multiple letters per year, use title for uniqueness
        date_text = doc.paragraphs[start].text.strip()
        safe_date = re.sub(r'\s+', '', date_text)
        filename = f"{year}_{safe_date}.md"
    else:
        filename = f"{year}.md"
    
    filepath = os.path.join(dir_path, filename)
    
    # If file already exists (multiple sections same year), append
    mode = 'a' if os.path.exists(filepath) else 'w'
    
    content = get_text_with_tables(doc, start, end)
    
    if not content.strip():
        return None
    
    with open(filepath, mode, encoding='utf-8') as f:
        if mode == 'w':
            f.write(f"# {section['title']}\n\n")
        else:
            f.write(f"\n\n---\n\n# {section['title']}\n\n")
        f.write(content)
        f.write('\n')
    
    return filepath


def write_special_sections(doc, output_base, first_main_section_idx):
    """Write special sections: preface, biography, appendix."""
    
    # 1. Preface (前言) - from start to biography
    bio_start = None
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if '巴菲特 40 岁' in text or '1930-1970' in text:
            bio_start = i
            break
    
    if bio_start:
        content = get_text_with_tables(doc, 0, bio_start)
        filepath = os.path.join(output_base, '前言.md')
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write("# 前言\n\n")
            f.write(content)
            f.write('\n')
    
    # 2. Biography (生平) - from bio_start to first early article
    early_start = 129  # "我最看好的股票"
    if bio_start:
        content = get_text_with_tables(doc, bio_start, early_start)
        filepath = os.path.join(output_base, '生平', '巴菲特生平_1930-1970.md')
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write("# 巴菲特生平 (1930-1970)\n\n")
            f.write(content)
            f.write('\n')
    
    # 3. Appendix (附录) - 道指百年走势, near end of doc
    appendix_start = None
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if '道指百年走势' in text:
            appendix_start = i
            break
    
    if appendix_start:
        content = get_text_with_tables(doc, appendix_start, len(doc.paragraphs))
        filepath = os.path.join(output_base, '附录', '道指百年走势.md')
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write("# 道指百年走势\n\n")
            f.write(content)
            f.write('\n')


def build_index(output_base):
    """Build a master index file."""
    index_lines = [
        "# 巴菲特文集 · 总目录",
        "",
        "> 巴菲特1950-2022致股东信 + 股东大会实录",
        "> 约440万字 | 覆盖72年投资智慧",
        "",
        "---",
        ""
    ]
    
    # Walk through directories
    categories = [
        ('前言.md', '📖 前言'),
        ('生平', '👤 生平'),
        ('早期文章', '📝 早期文章 (1951-1953)'),
        ('合伙人信', '💼 合伙人信 (1957-1969)'),
        ('致股东信', '📮 致股东信 (1965-2022)'),
        ('股东大会', '🎤 股东大会实录 (1986-2023)'),
        ('附录', '📊 附录'),
    ]
    
    for path, title in categories:
        full_path = os.path.join(output_base, path)
        
        if os.path.isfile(full_path):
            rel = os.path.relpath(full_path, output_base)
            index_lines.append(f"## {title}")
            index_lines.append(f"- [{os.path.basename(path)}]({rel})")
            index_lines.append("")
            continue
        
        if not os.path.isdir(full_path):
            continue
        
        files = sorted(os.listdir(full_path))
        md_files = [f for f in files if f.endswith('.md')]
        
        if not md_files:
            continue
        
        index_lines.append(f"## {title}")
        index_lines.append("")
        
        for f in md_files:
            rel = os.path.relpath(os.path.join(full_path, f), output_base)
            name = f.replace('.md', '')
            index_lines.append(f"- [{name}]({rel})")
        
        index_lines.append("")
    
    # Stats
    total_files = 0
    for root, dirs, files in os.walk(output_base):
        total_files += len([f for f in files if f.endswith('.md')])
    
    index_lines.extend([
        "---",
        "",
        f"共 {total_files} 个文件",
        "",
        "*由 WorkBuddy 自动生成*"
    ])
    
    index_path = os.path.join(output_base, 'README.md')
    with open(index_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(index_lines))
    
    return index_path, total_files


def main():
    print("正在加载文档...")
    doc = Document(DOC_PATH)
    print(f"文档加载完成: {len(doc.paragraphs)} 段落, {len(doc.tables)} 表格")
    
    # Find all sections
    print("\n正在识别章节...")
    main_sections = find_sections(doc)
    partnership_sections = find_partnership_letters(doc)
    early_sections = find_early_articles(doc)
    
    print(f"  致股东信/股东大会: {len(main_sections)} 个")
    print(f"  合伙人信: {len(partnership_sections)} 个")
    print(f"  早期文章: {len(early_sections)} 个")
    
    # Merge all
    all_sections = main_sections + partnership_sections + early_sections
    all_sections = merge_and_sort_sections(all_sections)
    print(f"  合并去重后: {len(all_sections)} 个章节")
    
    # Write special sections first
    print("\n正在写入特殊章节...")
    first_idx = all_sections[0]['para_idx'] if all_sections else len(doc.paragraphs)
    write_special_sections(doc, OUTPUT_BASE, first_idx)
    
    # Write each section
    print("正在拆分并写入各章节...")
    written = 0
    for idx, section in enumerate(all_sections):
        # Next section's start is this section's end
        if idx + 1 < len(all_sections):
            next_idx = all_sections[idx + 1]['para_idx']
        else:
            # Last section - go to end (but before appendix if exists)
            next_idx = len(doc.paragraphs)
        
        result = write_section(doc, section, next_idx, OUTPUT_BASE)
        if result:
            written += 1
            if written % 10 == 0:
                print(f"  已写入 {written} 个文件...")
    
    print(f"  共写入 {written} 个章节文件")
    
    # Build index
    print("\n正在生成总目录...")
    index_path, total = build_index(OUTPUT_BASE)
    print(f"总目录已生成: {index_path}")
    print(f"共 {total} 个Markdown文件")
    
    # Print summary
    print("\n=== 拆分完成 ===")
    for cat in ['致股东信', '股东大会', '合伙人信', '早期文章', '附录', '生平']:
        cat_dir = os.path.join(OUTPUT_BASE, cat)
        if os.path.isdir(cat_dir):
            count = len([f for f in os.listdir(cat_dir) if f.endswith('.md')])
            print(f"  {cat}: {count} 个文件")


if __name__ == '__main__':
    main()
